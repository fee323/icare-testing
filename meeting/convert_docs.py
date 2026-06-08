import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MEETING_DIR = r"D:\ICare_app-wajahat\meeting"
GREEN = RGBColor(0x00, 0x80, 0x00)
RED   = RGBColor(0xCC, 0x00, 0x00)

SPAN_RE = re.compile(
    r'<span style="color:(green|red)">\*\*\[(DONE|PENDING)\]\*\*</span>'
)

def clean_md(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*',     r'\1', text)
    text = re.sub(r'`(.*?)`',       r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    text = re.sub(r'<!--.*?-->', '', text)
    return text

def add_inline(para, text, base_bold=False):
    parts = SPAN_RE.split(text)
    idx = 0
    while idx < len(parts):
        chunk = parts[idx]
        if idx + 2 < len(parts) and parts[idx+1] in ('green', 'red'):
            color_name = parts[idx+1]
            label      = parts[idx+2]
            if chunk:
                run = para.add_run(clean_md(chunk))
                run.bold = base_bold
            run = para.add_run(f'[{label}]')
            run.bold = True
            run.font.color.rgb = GREEN if color_name == 'green' else RED
            idx += 3
        else:
            if chunk:
                run = para.add_run(clean_md(chunk))
                run.bold = base_bold
            idx += 1

def is_table_row(line):
    s = line.strip()
    return s.startswith('|') and s.endswith('|')

def is_sep_row(line):
    return is_table_row(line) and re.match(r'^[\|\s\-:]+$', line.strip())

def parse_cells(line):
    parts = line.strip().split('|')
    return [p.strip() for p in parts if p.strip()]

def convert(md_path, docx_path):
    doc = Document()
    # Set default font
    for style_name in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3'):
        try:
            doc.styles[style_name].font.name = 'Calibri'
        except Exception:
            pass

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        raw  = lines[i].rstrip('\n')
        line = raw.strip()
        i += 1

        if not line:
            continue

        # ---------- Headings ----------
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = min(len(m.group(1)), 4)
            text  = m.group(2)
            p = doc.add_heading('', level=level)
            p.clear()
            add_inline(p, text, base_bold=True)
            continue

        # ---------- Horizontal rule ----------
        if re.match(r'^[-_*]{3,}$', line):
            doc.add_paragraph('─' * 60)
            continue

        # ---------- Table ----------
        if is_table_row(line):
            table_lines = [line]
            while i < len(lines) and is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            data_rows = [r for r in table_lines if not is_sep_row(r)]
            if not data_rows:
                continue
            cols = len(parse_cells(data_rows[0]))
            table = doc.add_table(rows=0, cols=cols)
            table.style = 'Table Grid'
            for ri, row_line in enumerate(data_rows):
                cells = parse_cells(row_line)
                row = table.add_row()
                for ci in range(cols):
                    cell_text = cells[ci] if ci < len(cells) else ''
                    p = row.cells[ci].paragraphs[0]
                    p.clear()
                    add_inline(p, cell_text, base_bold=(ri == 0))
            doc.add_paragraph()
            continue

        # ---------- Bullet ----------
        bm = re.match(r'^[-•*]\s+(.*)', line)
        if bm:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, bm.group(1))
            continue

        # ---------- Numbered list ----------
        nm = re.match(r'^\d+[.)]\s+(.*)', line)
        if nm:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, nm.group(1))
            continue

        # ---------- Regular paragraph ----------
        p = doc.add_paragraph()
        add_inline(p, line)

    doc.save(docx_path)
    return True

errors = []
count  = 0
for fname in sorted(os.listdir(MEETING_DIR)):
    if not fname.endswith('.md'):
        continue
    src = os.path.join(MEETING_DIR, fname)
    dst = os.path.join(MEETING_DIR, fname[:-3] + '.docx')
    try:
        convert(src, dst)
        print(f"OK  {fname}")
        count += 1
    except Exception as e:
        print(f"ERR {fname}: {e}")
        errors.append(fname)

print(f"\n{count} files converted. {len(errors)} errors.")
