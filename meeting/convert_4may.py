import re, os
from docx import Document
from docx.shared import RGBColor

MEETING_DIR = r"D:\ICare_app-wajahat\meeting"
GREEN = RGBColor(0x00, 0x80, 0x00)
RED   = RGBColor(0xCC, 0x00, 0x00)
SPAN_RE = re.compile(r'<span style="color:(green|red)">\*\*\[(DONE|PENDING)\]\*\*</span>')

def clean_md(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text

def add_inline(para, text, base_bold=False):
    parts = SPAN_RE.split(text)
    idx = 0
    while idx < len(parts):
        chunk = parts[idx]
        if idx + 2 < len(parts) and parts[idx+1] in ('green', 'red'):
            if chunk:
                r = para.add_run(clean_md(chunk)); r.bold = base_bold
            r = para.add_run(f'[{parts[idx+2]}]'); r.bold = True
            r.font.color.rgb = GREEN if parts[idx+1] == 'green' else RED
            idx += 3
        else:
            if chunk:
                r = para.add_run(clean_md(chunk)); r.bold = base_bold
            idx += 1

def is_table_row(l):
    s = l.strip()
    return s.startswith('|') and s.endswith('|')

def is_sep(l):
    return is_table_row(l) and re.match(r'^[\|\s\-:]+$', l.strip())

def parse_cells(l):
    return [p.strip() for p in l.strip().split('|') if p.strip()]

src = os.path.join(MEETING_DIR, '4 may docs.md')
dst = os.path.join(MEETING_DIR, '4 may docs.docx')

doc = Document()
with open(src, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    raw = lines[i].rstrip('\n')
    line = raw.strip()
    i += 1
    if not line:
        continue
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        p = doc.add_heading('', level=min(len(m.group(1)), 4))
        p.clear()
        add_inline(p, m.group(2), True)
        continue
    if is_table_row(line):
        tl = [line]
        while i < len(lines) and is_table_row(lines[i].strip()):
            tl.append(lines[i].strip())
            i += 1
        dr = [r for r in tl if not is_sep(r)]
        if not dr:
            continue
        cols = len(parse_cells(dr[0]))
        table = doc.add_table(rows=0, cols=cols)
        table.style = 'Table Grid'
        for ri, rl in enumerate(dr):
            cs = parse_cells(rl)
            row = table.add_row()
            for ci in range(cols):
                ct = cs[ci] if ci < len(cs) else ''
                p = row.cells[ci].paragraphs[0]
                p.clear()
                add_inline(p, ct, ri == 0)
        doc.add_paragraph()
        continue
    bm = re.match(r'^[-•*]\s+(.*)', line)
    if bm:
        p = doc.add_paragraph(style='List Bullet')
        add_inline(p, bm.group(1))
        continue
    nm = re.match(r'^\d+[.)]\s+(.*)', line)
    if nm:
        p = doc.add_paragraph(style='List Number')
        add_inline(p, nm.group(1))
        continue
    p = doc.add_paragraph()
    add_inline(p, line)

doc.save(dst)
print("Done: 4 may docs.docx")
